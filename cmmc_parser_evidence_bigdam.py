#!/usr/bin/env python3
"""
CMMC SSP Parser with Evidence Enrichment
Integrates evidence CSV to enrich Evidence_Strings with links and descriptions
"""

import pandas as pd
import json
import os
import sys
import re
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import argparse
import logging

# Control family mappings
CONTROL_FAMILIES = {
    '3.1': 'AC - Access Control',
    '3.2': 'AT - Awareness and Training', 
    '3.3': 'AU - Audit and Accountability',
    '3.4': 'CM - Configuration Management',
    '3.5': 'IA - Identification and Authentication',
    '3.6': 'IR - Incident Response',
    '3.7': 'MA - Maintenance',
    '3.8': 'MP - Media Protection',
    '3.9': 'PS - Personnel Security',
    '3.10': 'PE - Physical Protection',
    '3.11': 'RA - Risk Assessment',
    '3.12': 'SA - Security Assessment',
    '3.13': 'SC - System and Communications Protection',
    '3.14': 'SI - System and Information Integrity'
}

class CMMCParser:
    def __init__(self, config_file='config.json'):
        """Initialize parser with configuration"""
        self.config = self.load_config(config_file)
        self.validation_errors = []
        self.validation_warnings = []
        self.evidence_map = {}  # Will store enriched evidence
        self.setup_logging()
        
    def load_config(self, config_file):
        """Load configuration from JSON file"""
        if os.path.exists(config_file):
            with open(config_file, 'r') as f:
                return json.load(f)
        else:
            # Default configuration
            return {
                "input_csv": "ssp_prime_to_csv.csv",
                "evidence_csv": "evidence_enrichment.csv",  # New field
                "output_dir": "./output",
                "generate_html": True,
                "generate_docx": True,
                "validate_poam_rules": True,
                "evidence_base_path": "/CMMC_Evidence/"
            }
    
    def setup_logging(self):
        """Setup logging configuration"""
        os.makedirs(self.config['output_dir'], exist_ok=True)
        log_file = os.path.join(self.config['output_dir'], f'parser_log_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')
        
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)
    
    def load_evidence_enrichment(self, evidence_csv_path):
        """Load and process the evidence enrichment CSV"""
        if not os.path.exists(evidence_csv_path):
            self.logger.warning(f"Evidence enrichment file not found: {evidence_csv_path}")
            return
        
        self.logger.info(f"Loading evidence enrichment from: {evidence_csv_path}")
        
        # Detect delimiter
        with open(evidence_csv_path, 'r', encoding='utf-8-sig') as f:
            first_line = f.readline()
            if '|' in first_line and first_line.count('|') > first_line.count(','):
                delimiter = '|'
            else:
                delimiter = ','
        
        # Load the evidence CSV
        evidence_df = pd.read_csv(evidence_csv_path, delimiter=delimiter, encoding='utf-8-sig')
        self.logger.info(f"Evidence CSV columns: {list(evidence_df.columns)}")
        
        # Process each evidence row
        for index, row in evidence_df.iterrows():
            # Skip rows with no mappings
            if pd.isna(row.get('Suggested_CMMC_Mappings', '')) and pd.isna(row.get('Provided_CMMC_Mappings', '')):
                continue
            
            # Skip rows marked as IGNORE
            if 'Description' in row and not pd.isna(row['Description']):
                if 'IGNORE' in str(row['Description']).upper():
                    self.logger.info(f"Skipping ignored evidence: {row.get('File_Name', '')}")
                    continue
            
            # Format the evidence entry
            evidence_entry = self.format_evidence_entry(row)
            
            # Get CMMC mappings - prioritize Suggested_CMMC_Mappings
            cmmc_mappings = []
            
            # Check Suggested_CMMC_Mappings first (PRIMARY)
            if 'Suggested_CMMC_Mappings' in row and not pd.isna(row['Suggested_CMMC_Mappings']) and row['Suggested_CMMC_Mappings'] != '':
                suggested = str(row['Suggested_CMMC_Mappings'])
                cmmc_mappings = self.parse_delimited_content(suggested)
            # ONLY use Provided_CMMC_Mappings if Suggested is empty (FALLBACK)
            elif 'Provided_CMMC_Mappings' in row and not pd.isna(row['Provided_CMMC_Mappings']) and row['Provided_CMMC_Mappings'] != '':
                provided = str(row['Provided_CMMC_Mappings'])
                cmmc_mappings = self.parse_delimited_content(provided)
                self.logger.debug(f"Using Provided_CMMC_Mappings as fallback for {row.get('File_Name', '')}")
            
            # Add evidence to all mapped controls
            for cmmc_id in cmmc_mappings:
                cmmc_id = cmmc_id.strip()
                if cmmc_id:
                    if cmmc_id not in self.evidence_map:
                        self.evidence_map[cmmc_id] = []
                    self.evidence_map[cmmc_id].append(evidence_entry)
                    self.logger.debug(f"Added evidence to control {cmmc_id}: {evidence_entry[:50]}...")
        
        self.logger.info(f"Loaded evidence for {len(self.evidence_map)} controls")
        
        # Log summary of evidence per control
        for cmmc_id, evidence_list in self.evidence_map.items():
            self.logger.debug(f"Control {cmmc_id}: {len(evidence_list)} evidence items")
    
    def format_evidence_entry(self, row):
        """Format an evidence entry with link and description"""
        entry_parts = []
        
        # Add SharePoint link if present
        if 'Current_Sharepoint_Link' in row and not pd.isna(row['Current_Sharepoint_Link']) and row['Current_Sharepoint_Link'] != '':
            link = str(row['Current_Sharepoint_Link']).strip()
            entry_parts.append(f"[{link}]")
        
        # Add description if present
        if 'Description' in row and not pd.isna(row['Description']) and row['Description'] != '':
            description = str(row['Description']).strip()
            # Check if description should be ignored
            if 'IGNORE' not in description.upper():
                entry_parts.append(description)
        
        # If no description, use File_Name as fallback
        elif 'File_Name' in row and not pd.isna(row['File_Name']) and row['File_Name'] != '':
            file_name = str(row['File_Name']).strip()
            entry_parts.append(file_name)
        
        # Join with space-dash-space if we have both link and description
        if len(entry_parts) == 2:
            return f"{entry_parts[0]} - {entry_parts[1]}"
        elif entry_parts:
            return entry_parts[0]
        else:
            return ""
    
    def parse_delimited_content(self, content):
        """Parse pipe or semicolon-delimited content into bullet points"""
        if pd.isna(content) or content == '':
            return []
        
        # Convert to string
        content = str(content)
        items = []
        
        # Replace pipes with semicolons for uniform processing
        content = content.replace('|', ';')
        
        # Split by semicolon
        for item in content.split(';'):
            item = item.strip()
            # Skip empty items and placeholder text
            if item and item not in ['', 'header', 'bullet_1', 'bullet_2', 'bullet1', 'bullet2']:
                # Remove trailing periods for consistency
                if item.endswith('.'):
                    item = item[:-1].strip()
                items.append(item)
        
        return items
    
    def get_enriched_evidence_strings(self, cmmc_id, existing_evidence):
        """Combine existing evidence strings with enriched evidence"""
        evidence_items = []
        
        # Parse existing evidence if present
        if existing_evidence and not pd.isna(existing_evidence) and existing_evidence != '':
            evidence_items.extend(self.parse_delimited_content(existing_evidence))
        
        # Add enriched evidence if available
        if cmmc_id in self.evidence_map:
            evidence_items.extend(self.evidence_map[cmmc_id])
        
        # Remove duplicates while preserving order
        seen = set()
        unique_items = []
        for item in evidence_items:
            if item not in seen:
                seen.add(item)
                unique_items.append(item)
        
        return unique_items
    
    def get_control_family(self, cmmc_id):
        """Determine control family from CMMC ID"""
        cmmc_id = str(cmmc_id)
        prefix = '.'.join(cmmc_id.split('.')[:2])
        return CONTROL_FAMILIES.get(prefix, 'Unknown')
    
    def get_control_title(self, control_text):
        """Extract the full control title from control text"""
        if pd.isna(control_text) or control_text == '':
            return "Untitled Control"
        # Parse pipes as list items and join with commas for the title
        text = str(control_text)
        items = self.parse_delimited_content(text)
        if items:
            # Join the parsed items with commas for a readable title
            text = ', '.join(items)
        # Remove trailing period for consistency
        if text.endswith('.'):
            text = text[:-1]
        return text
    
    def validate_csv_data(self, df):
        """Validate CSV data according to business rules"""
        self.logger.info("Starting data validation...")
        
        for index, row in df.iterrows():
            # Check for CMMC_ID
            if pd.isna(row['CMMC_ID']) or str(row['CMMC_ID']).strip() == '':
                self.validation_errors.append(f"Row {index}: Missing CMMC_ID")
                continue
            
            cmmc_id = str(row['CMMC_ID'])
            
            if pd.isna(row['Control']) or row['Control'] == '':
                self.validation_warnings.append(f"Control {cmmc_id}: Missing Control description")
            
            # Validate Score values
            if row['Score'] not in [1, 3, 5]:
                self.validation_errors.append(f"Control {cmmc_id}: Invalid Score value '{row['Score']}' (must be 1, 3, or 5)")
            
            # POA&M validation - only Score=1 can be POA&M'd
            if row['Score'] in [3, 5] and str(row['AR_CAP_POAM']).upper() in ['POA&M', 'POAM']:
                self.validation_errors.append(
                    f"Control {cmmc_id}: CRITICAL - Score={row['Score']} (non-POAMable) but AR_CAP_POAM='{row['AR_CAP_POAM']}'"
                )
        
        self.logger.info(f"Validation complete: {len(self.validation_errors)} errors, {len(self.validation_warnings)} warnings")
    
    def generate_html_for_control(self, row):
        """Generate HTML structure for a single control"""
        cmmc_id = str(row['CMMC_ID'])
        control_title = self.get_control_title(row['Control'])
        
        # Determine implementation status
        if str(row['AR_CAP_POAM']).upper() in ['POA&M', 'POAM']:
            impl_status = "Plan of Action & Milestones (POA&M)"
        elif str(row['AR_CAP_POAM']).upper() == 'AUDIT READY':
            impl_status = "Audit Ready"
        else:
            impl_status = "Implemented"
        
        html = f"""
        <div class="control-section" id="{cmmc_id.replace('.', '_')}">
            <h2>{cmmc_id} - {control_title}</h2>
            
            <div class="practice-statement">
                <h3>Practice Statement</h3>
                <p>{control_title}</p>
            </div>
            
            <div class="implementation-status">
                <h3>Implementation Status</h3>
                <p>{impl_status}</p>
            </div>
        """
        
        # Add Policy Summary if present
        if not pd.isna(row.get('Policy_Statement', '')) and row.get('Policy_Statement', '') != '':
            policy_content = str(row['Policy_Statement'])
            if policy_content not in ['header;bullet_1;bullet_2', '']:
                html += """
                <div class="policy-summary">
                    <h3>Policy Summary</h3>
                """
                policy_items = self.parse_delimited_content(policy_content)
                if policy_items:
                    html += "<ul>"
                    for item in policy_items:
                        html += f"<li>{item}</li>"
                    html += "</ul>"
                else:
                    html += f"<p>{policy_content}</p>"
                html += "</div>"
        
        # Add Implementation Details
        html += """
            <div class="implementation-details">
                <h3>Implementation Details</h3>
        """
        
        # Azure Mechanism
        if 'Azure_Mechanism' in row and not pd.isna(row['Azure_Mechanism']) and row['Azure_Mechanism'] != '':
            html += """
                <div class="azure-mechanism">
                    <h4>Azure Mechanism</h4>
            """
            azure_items = self.parse_delimited_content(row['Azure_Mechanism'])
            if len(azure_items) > 1:
                html += "<ul>"
                for item in azure_items:
                    html += f"<li>{item}</li>"
                html += "</ul>"
            elif azure_items:
                html += f"<p>{azure_items[0]}</p>"
            else:
                html += f"<p>{row['Azure_Mechanism']}</p>"
            html += "</div>"
        
        # Azure Configuration Process
        if 'Azure_Configuration_Process' in row and not pd.isna(row['Azure_Configuration_Process']) and row['Azure_Configuration_Process'] != '':
            html += """
                <div class="azure-config">
                    <h4>Azure Configuration Process</h4>
            """
            config_items = self.parse_delimited_content(row['Azure_Configuration_Process'])
            if len(config_items) > 1:
                html += "<ul>"
                for item in config_items:
                    html += f"<li>{item}</li>"
                html += "</ul>"
            elif config_items:
                html += f"<p>{config_items[0]}</p>"
            else:
                html += f"<p>{row['Azure_Configuration_Process']}</p>"
            html += "</div>"
        
        # Azure Evidence
        if 'Azure_Evidence' in row and not pd.isna(row['Azure_Evidence']) and row['Azure_Evidence'] != '':
            html += """
                <div class="azure-evidence">
                    <h4>Azure Evidence</h4>
                    <ul>
            """
            for evidence in self.parse_delimited_content(row['Azure_Evidence']):
                html += f"<li>{evidence}</li>"
            html += """
                    </ul>
                </div>
            """
        
        # AVD/Laptop
        if 'AVD_Laptop' in row and not pd.isna(row['AVD_Laptop']) and row['AVD_Laptop'] != '':
            html += """
                <div class="avd-laptop">
                    <h4>AVD/Laptop Environment</h4>
            """
            avd_items = self.parse_delimited_content(row['AVD_Laptop'])
            if len(avd_items) > 1:
                html += "<ul>"
                for item in avd_items:
                    html += f"<li>{item}</li>"
                html += "</ul>"
            elif avd_items:
                html += f"<p>{avd_items[0]}</p>"
            else:
                html += f"<p>{row['AVD_Laptop']}</p>"
            
            # AVD Evidence
            if 'AVD_Laptop_Evidence' in row and not pd.isna(row['AVD_Laptop_Evidence']) and row['AVD_Laptop_Evidence'] != '':
                html += "<p><strong>Evidence:</strong></p><ul>"
                for evidence in self.parse_delimited_content(row['AVD_Laptop_Evidence']):
                    html += f"<li>{evidence}</li>"
                html += "</ul>"
            html += "</div>"
        
        # Evidence_Strings (enriched with evidence CSV data)
        evidence_items = self.get_enriched_evidence_strings(cmmc_id, row.get('Evidence_Strings', ''))
        if evidence_items:
            html += """
                <div class="evidence-strings">
                    <h4>Evidence Strings</h4>
                    <ul>
            """
            for evidence in evidence_items:
                html += f"<li>{evidence}</li>"
            html += """
                    </ul>
                </div>
            """
        
        html += "</div>"  # Close implementation-details
        html += "</div>"  # Close control-section
        
        return html
    
    def generate_docx_for_control(self, doc, control):
        """Add a control to the DOCX document"""
        # Control heading
        control_title = self.get_control_title(control['Control'])
        doc.add_heading(f'{control["CMMC_ID"]} - {control_title}', 1)
        
        # Practice Statement
        doc.add_heading('Practice Statement', 2)
        doc.add_paragraph(control_title)
        
        # Implementation Status
        doc.add_heading('Implementation Status', 2)
        if str(control['AR_CAP_POAM']).upper() in ['POA&M', 'POAM']:
            status = "Plan of Action & Milestones (POA&M)"
        elif str(control['AR_CAP_POAM']).upper() == 'AUDIT READY':
            status = "Audit Ready"
        else:
            status = "Implemented"
        doc.add_paragraph(status)
        
        # Policy Summary (if present)
        if not pd.isna(control.get('Policy_Statement', '')) and control.get('Policy_Statement', '') != '':
            policy_content = str(control['Policy_Statement'])
            if policy_content not in ['header;bullet_1;bullet_2', '']:
                doc.add_heading('Policy Summary', 2)
                policy_items = self.parse_delimited_content(policy_content)
                if policy_items:
                    for item in policy_items:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(item)
                else:
                    doc.add_paragraph(policy_content)
        
        # Implementation Details
        doc.add_heading('Implementation Details', 2)
        
        # Azure Mechanism
        if 'Azure_Mechanism' in control and not pd.isna(control['Azure_Mechanism']) and control['Azure_Mechanism'] != '':
            doc.add_heading('Azure Mechanism', 3)
            azure_items = self.parse_delimited_content(control['Azure_Mechanism'])
            if len(azure_items) > 1:
                for item in azure_items:
                    if item:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(item)
            elif azure_items:
                doc.add_paragraph(azure_items[0])
            else:
                doc.add_paragraph(control['Azure_Mechanism'])
        
        # Azure Configuration Process
        if 'Azure_Configuration_Process' in control and not pd.isna(control['Azure_Configuration_Process']) and control['Azure_Configuration_Process'] != '':
            doc.add_heading('Azure Configuration Process', 3)
            config_items = self.parse_delimited_content(control['Azure_Configuration_Process'])
            if len(config_items) > 1:
                for item in config_items:
                    if item:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(item)
            elif config_items:
                doc.add_paragraph(config_items[0])
            else:
                doc.add_paragraph(control['Azure_Configuration_Process'])
        
        # Azure Evidence
        if 'Azure_Evidence' in control and not pd.isna(control['Azure_Evidence']) and control['Azure_Evidence'] != '':
            doc.add_heading('Azure Evidence', 3)
            for evidence in self.parse_delimited_content(control['Azure_Evidence']):
                if evidence:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(evidence)
        
        # AVD/Laptop
        if 'AVD_Laptop' in control and not pd.isna(control['AVD_Laptop']) and control['AVD_Laptop'] != '':
            doc.add_heading('AVD/Laptop Environment', 3)
            avd_items = self.parse_delimited_content(control['AVD_Laptop'])
            if len(avd_items) > 1:
                for item in avd_items:
                    if item:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(item)
            elif avd_items:
                doc.add_paragraph(avd_items[0])
            else:
                doc.add_paragraph(control['AVD_Laptop'])
            
            # AVD Evidence
            if 'AVD_Laptop_Evidence' in control and not pd.isna(control['AVD_Laptop_Evidence']) and control['AVD_Laptop_Evidence'] != '':
                p = doc.add_paragraph()
                p.add_run('Evidence:').bold = True
                for evidence in self.parse_delimited_content(control['AVD_Laptop_Evidence']):
                    if evidence:
                        doc.add_paragraph(evidence, style='List Bullet')
        
        # Evidence_Strings (enriched with evidence CSV data)
        cmmc_id = str(control['CMMC_ID'])
        evidence_items = self.get_enriched_evidence_strings(cmmc_id, control.get('Evidence_Strings', ''))
        if evidence_items:
            doc.add_heading('Evidence Strings', 3)
            for evidence in evidence_items:
                if evidence:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(evidence)
    
    def generate_html_files(self, df):
        """Generate HTML files grouped by control family"""
        self.logger.info("Generating HTML files...")
        
        # Group controls by family
        families = {}
        for index, row in df.iterrows():
            if pd.isna(row['CMMC_ID']):
                continue
            family = self.get_control_family(row['CMMC_ID'])
            family_key = family.split(' - ')[0]
            
            if family_key not in families:
                families[family_key] = []
            families[family_key].append(row)
        
        # Generate HTML for each family
        for family_key, controls in families.items():
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CMMC 2.0 SSP - {family_key}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #003366; border-bottom: 3px solid #003366; padding-bottom: 10px; }}
        h2 {{ color: #0066cc; border-bottom: 1px solid #0066cc; padding-bottom: 5px; }}
        h3 {{ color: #333333; background-color: #f0f0f0; padding: 5px; }}
        h4 {{ color: #666666; margin-top: 15px; }}
        .control-section {{ margin-bottom: 40px; page-break-after: always; }}
        .practice-statement, .implementation-status, .policy-summary, 
        .implementation-details, .evidence-strings {{ margin: 20px 0; }}
        ul {{ margin-left: 20px; }}
        li {{ margin: 5px 0; }}
        strong {{ color: #000080; }}
    </style>
</head>
<body>
    <h1>{CONTROL_FAMILIES.get('.'.join(controls[0]['CMMC_ID'].split('.')[:2]), family_key)}</h1>
    <p>Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
"""
            
            for control in controls:
                html_content += self.generate_html_for_control(control)
            
            html_content += """
</body>
</html>
"""
            
            output_path = os.path.join(self.config['output_dir'], f'{family_key}_controls.html')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"Generated HTML: {output_path}")
    
    def generate_docx_files(self, df):
        """Generate DOCX files grouped by control family"""
        self.logger.info("Generating DOCX files...")
        
        # Group controls by family
        families = {}
        for index, row in df.iterrows():
            if pd.isna(row['CMMC_ID']):
                continue
            family = self.get_control_family(row['CMMC_ID'])
            family_key = family.split(' - ')[0]
            
            if family_key not in families:
                families[family_key] = []
            families[family_key].append(row)
        
        # Generate DOCX for each family
        for family_key, controls in families.items():
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'CMMC Level 2 - {CONTROL_FAMILIES.get(".".join(controls[0]["CMMC_ID"].split(".")[:2]), family_key)}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add generation date
            date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process each control
            for i, control in enumerate(controls):
                if i > 0:
                    doc.add_page_break()
                self.generate_docx_for_control(doc, control)
            
            # Display ASCII art after DOCX completes
            if family_key == list(families.keys())[-1]:  # Only show after last family
                print("\n")
                print("╔══════════════════════════════════════════════════════════╗")
                print("║  YOU HAVE BEEN BIG DAM JUDGED                           ║")
                print("╚══════════════════════════════════════════════════════════╝")
                print("\n")
            
            # Save DOCX file
            output_path = os.path.join(self.config['output_dir'], f'{family_key}_controls.docx')
            doc.save(output_path)
            
            self.logger.info(f"Generated DOCX: {output_path}")
    
    def generate_validation_report(self, df):
        """Generate validation report"""
        report_path = os.path.join(self.config['output_dir'], 'validation_report.txt')
        
        with open(report_path, 'w') as f:
            f.write("="*80 + "\n")
            f.write("CMMC SSP Parser Validation Report\n")
            f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("="*80 + "\n\n")
            
            f.write("SUMMARY STATISTICS\n")
            f.write("-"*40 + "\n")
            f.write(f"Total Controls Processed: {len(df)}\n")
            
            # Evidence enrichment summary
            if self.evidence_map:
                f.write(f"\nEvidence Enrichment:\n")
                f.write(f"  Controls with enriched evidence: {len(self.evidence_map)}\n")
                total_evidence = sum(len(v) for v in self.evidence_map.values())
                f.write(f"  Total evidence entries: {total_evidence}\n")
            
            # POA&M Summary
            f.write("\nPOA&M Summary:\n")
            poam_count = len(df[df['AR_CAP_POAM'].astype(str).str.upper().isin(['POA&M', 'POAM'])])
            audit_ready = len(df[df['AR_CAP_POAM'].astype(str).str.upper() == 'AUDIT READY'])
            implemented = len(df) - poam_count - audit_ready
            f.write(f"  Implemented: {implemented}\n")
            f.write(f"  Audit Ready: {audit_ready}\n")
            f.write(f"  POA&M: {poam_count}\n")
            
            # Critical Errors
            if self.validation_errors:
                f.write("\nCRITICAL ERRORS (Must Fix)\n")
                f.write("-"*40 + "\n")
                for error in self.validation_errors:
                    f.write(f"❌ {error}\n")
            
            # Warnings
            if self.validation_warnings:
                f.write("\nWARNINGS (Should Review)\n")
                f.write("-"*40 + "\n")
                for warning in self.validation_warnings:
                    f.write(f"⚠️  {warning}\n")
        
        self.logger.info(f"Validation report saved: {report_path}")
    
    def filter_dataframe(self, df, controls=None, families=None, control_range=None):
        """Filter dataframe based on specific controls, families, or range"""
        if not controls and not families and not control_range:
            return df
        
        filtered_df = pd.DataFrame()
        
        if controls:
            self.logger.info(f"Filtering for controls: {controls}")
            control_filter = df['CMMC_ID'].astype(str).isin(controls)
            filtered_df = pd.concat([filtered_df, df[control_filter]])
        
        if families:
            self.logger.info(f"Filtering for families: {families}")
            family_dfs = []
            for index, row in df.iterrows():
                if pd.isna(row['CMMC_ID']):
                    continue
                family = self.get_control_family(row['CMMC_ID']).split(' - ')[0]
                if family in families:
                    family_dfs.append(row)
            if family_dfs:
                filtered_df = pd.concat([filtered_df, pd.DataFrame(family_dfs)])
        
        if control_range:
            self.logger.info(f"Filtering for range: {control_range}")
            try:
                start, end = control_range.split('-')
                start_parts = start.split('.')
                end_parts = end.split('.')
                
                range_dfs = []
                for index, row in df.iterrows():
                    if pd.isna(row['CMMC_ID']):
                        continue
                    control_id = str(row['CMMC_ID'])
                    control_parts = control_id.split('.')
                    
                    if (len(control_parts) == 3 and len(start_parts) == 3 and
                        control_parts[0] == start_parts[0] and 
                        control_parts[1] == start_parts[1] and
                        int(control_parts[2]) >= int(start_parts[2]) and
                        int(control_parts[2]) <= int(end_parts[2])):
                        range_dfs.append(row)
                
                if range_dfs:
                    filtered_df = pd.concat([filtered_df, pd.DataFrame(range_dfs)])
            except Exception as e:
                self.logger.error(f"Invalid range format: {control_range}")
        
        if not filtered_df.empty:
            filtered_df = filtered_df.drop_duplicates(subset=['CMMC_ID'])
            self.logger.info(f"Filtered to {len(filtered_df)} controls")
        
        return filtered_df
    
    def run(self, filter_controls=None, filter_families=None, filter_range=None):
        """Main execution method"""
        try:
            # Load evidence enrichment first
            if 'evidence_csv' in self.config:
                self.load_evidence_enrichment(self.config['evidence_csv'])
            
            # Load main CSV
            self.logger.info(f"Loading CSV: {self.config['input_csv']}")
            
            # Try to detect delimiter
            with open(self.config['input_csv'], 'r', encoding='utf-8-sig') as f:
                first_line = f.readline()
                if '|' in first_line and first_line.count('|') > first_line.count(','):
                    delimiter = '|'
                    self.logger.info("Detected pipe-delimited file")
                else:
                    delimiter = ','
                    self.logger.info("Using comma delimiter")
            
            df = pd.read_csv(self.config['input_csv'], delimiter=delimiter, encoding='utf-8-sig')
            
            # Log the columns found
            self.logger.info(f"Columns found in CSV: {list(df.columns)}")
            
            # Apply filters if specified
            if filter_controls or filter_families or filter_range:
                df = self.filter_dataframe(df, filter_controls, filter_families, filter_range)
                if df.empty:
                    self.logger.warning("No controls matched the filter criteria")
                    print("\n⚠️  No controls matched your filter criteria!")
                    return 0
            
            # Validate data
            if self.config.get('validate_poam_rules', True):
                self.validate_csv_data(df)
            
            # Generate outputs
            if self.config.get('generate_html', True):
                self.generate_html_files(df)
            
            if self.config.get('generate_docx', True):
                self.generate_docx_files(df)
            
            # Generate validation report
            self.generate_validation_report(df)
            
            # Print summary
            print("\n" + "="*50)
            print("✅ CMMC SSP Parser Completed Successfully!")
            print("="*50)
            print(f"📁 Output directory: {self.config['output_dir']}")
            print(f"📊 Errors: {len(self.validation_errors)}")
            print(f"⚠️  Warnings: {len(self.validation_warnings)}")
            if self.evidence_map:
                print(f"📎 Evidence enriched for {len(self.evidence_map)} controls")
            
            if self.validation_errors:
                print("\n❌ Critical errors found - review validation report!")
                return 1
            else:
                print("\n✅ No critical errors - ready for review!")
                return 0
                
        except Exception as e:
            self.logger.error(f"Fatal error: {str(e)}", exc_info=True)
            print(f"\n❌ Fatal error: {str(e)}")
            return 1

def main():
    """Main entry point"""
    parser = argparse.ArgumentParser(description='CMMC SSP Parser with Evidence Enrichment')
    parser.add_argument('-c', '--config', default='config.json', help='Configuration file path')
    parser.add_argument('-i', '--input', help='Override input CSV file')
    parser.add_argument('-e', '--evidence', help='Override evidence CSV file')
    parser.add_argument('-o', '--output', help='Override output directory')
    parser.add_argument('--html-only', action='store_true', help='Generate only HTML files')
    parser.add_argument('--docx-only', action='store_true', help='Generate only DOCX files')
    parser.add_argument('--skip-validation', action='store_true', help='Skip POA&M validation')
    parser.add_argument('--controls', nargs='+', help='Process specific controls')
    parser.add_argument('--families', nargs='+', help='Process specific families')
    parser.add_argument('--range', help='Process control range')
    
    args = parser.parse_args()
    
    # Create parser instance
    cmmc_parser = CMMCParser(args.config)
    
    # Override config with command line arguments
    if args.input:
        cmmc_parser.config['input_csv'] = args.input
    if args.evidence:
        cmmc_parser.config['evidence_csv'] = args.evidence
    if args.output:
        cmmc_parser.config['output_dir'] = args.output
    if args.html_only:
        cmmc_parser.config['generate_docx'] = False
    if args.docx_only:
        cmmc_parser.config['generate_html'] = False
    if args.skip_validation:
        cmmc_parser.config['validate_poam_rules'] = False
    
    # Run parser
    sys.exit(cmmc_parser.run(
        filter_controls=args.controls,
        filter_families=args.families,
        filter_range=args.range
    ))

if __name__ == "__main__":
    main()
